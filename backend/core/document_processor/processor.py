# File Path: insightgen/backend/core/document_processor/processor.py
import os
import logging
import fitz
import pandas as pd
import numpy as np
import openpyxl
from typing import Dict, Any, List, Optional
from pathlib import Path
import pytesseract
from PIL import Image
import easyocr
from typing import Optional, Dict, Any, List
import os
from pathlib import Path
import fitz  # PyMuPDF
import json
import docx
import xml.etree.ElementTree as ET
import torch
import io
import csv
import yaml
from langdetect import detect
import spacy
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.processors = {}
        self.reader = easyocr.Reader(['en'])
        self.nlp = spacy.load("en_core_web_sm")
        try:
            self.reader = easyocr.Reader(['en'])  # Initialize EasyOCR
        except Exception as e:
            logger.warning(f"EasyOCR initialization failed: {e}")
            self.reader = None
        
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            logger.warning(f"Spacy model loading failed: {e}")
            self.nlp = None

        self.processors = {
            "basic": {
                "initialized": True,
                "status": "ready"
            }
        }
        
    async def _process_image(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process image files with OCR"""
        result = {
            "content": "",
            "metadata": {},
            "text_blocks": [],
            "analysis": {}
        }

        try:
            # Load image
            image = Image.open(file_path)
            
            # Get image metadata
            result["metadata"] = {
                "size": image.size,
                "mode": image.mode,
                "format": image.format,
                "dpi": image.info.get('dpi', 'Unknown'),
                "filename": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path)
            }

            # Convert image to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Perform OCR with Tesseract
            try:
                tesseract_text = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                
                # Filter and process Tesseract results
                text_blocks = []
                for i in range(len(tesseract_text['text'])):
                    if int(tesseract_text['conf'][i]) > 50:  # Filter low-confidence results
                        text_blocks.append({
                            'text': tesseract_text['text'][i],
                            'confidence': float(tesseract_text['conf'][i]),
                            'bbox': {
                                'x': tesseract_text['left'][i],
                                'y': tesseract_text['top'][i],
                                'width': tesseract_text['width'][i],
                                'height': tesseract_text['height'][i]
                            },
                            'source': 'tesseract'
                        })
            except Exception as e:
                logger.warning(f"Tesseract OCR failed: {e}")
                text_blocks = []

            # Perform OCR with EasyOCR as backup/complement
            try:
                easyocr_result = self.reader.readtext(np.array(image))
                
                for bbox, text, conf in easyocr_result:
                    text_blocks.append({
                        'text': text,
                        'confidence': float(conf) * 100,
                        'bbox': {
                            'points': bbox,
                        },
                        'source': 'easyocr'
                    })
            except Exception as e:
                logger.warning(f"EasyOCR failed: {e}")

            # Combine all extracted text
            all_text = []
            for block in text_blocks:
                if block['text'] and block['text'].strip():
                    all_text.append(block['text'].strip())

            # Store results
            result["text_blocks"] = text_blocks
            result["content"] = " ".join(all_text)

            # Basic text analysis if content exists
            if result["content"]:
                try:
                    doc = self.nlp(result["content"])
                    result["analysis"] = {
                        "sentences": len(list(doc.sents)),
                        "words": len([token for token in doc if not token.is_punct]),
                        "entities": [
                            {
                                "text": ent.text,
                                "label": ent.label_,
                                "start": ent.start_char,
                                "end": ent.end_char
                            }
                            for ent in doc.ents
                        ]
                    }
                except Exception as e:
                    logger.warning(f"Text analysis failed: {e}")
                    result["analysis"] = {}

            return result

        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise
        
    async def _process_text(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process plain text files"""
        result = {
            "content": "",
            "metadata": {},
            "analysis": {}
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                result["content"] = content
                
                result["metadata"] = {
                    "filename": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "encoding": file.encoding
                }

                # Language detection
                try:
                    result["metadata"]["language"] = detect(content)
                except Exception as e:
                    logger.warning(f"Language detection failed: {e}")
                    result["metadata"]["language"] = "unknown"

                # Text analysis
                if self.nlp and content.strip():
                    try:
                        doc = self.nlp(content)
                        result["analysis"] = {
                            "sentences": len(list(doc.sents)),
                            "words": len([token for token in doc if not token.is_punct]),
                            "entities": [
                                {
                                    "text": ent.text,
                                    "label": ent.label_,
                                    "start": ent.start_char,
                                    "end": ent.end_char
                                }
                                for ent in doc.ents
                            ]
                        }
                    except Exception as e:
                        logger.warning(f"Text analysis failed: {e}")

            return result

        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            raise

    async def process_document(
        self,
        file_path: str,
        file_type: str,
        extraction_type: str = "text",
        options: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Process document based on type and extraction needs"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        processors = {
            "pdf": self._process_pdf,
            "jpg": self._process_image,
            "jpeg": self._process_image,
            "png": self._process_image,
            "docx": self._process_docx,
            "doc": self._process_docx,
            "txt": self._process_text,
            "csv": self._process_csv,
            "json": self._process_json,
            "xml": self._process_xml,
            "yaml": self._process_yaml,
            "yml": self._process_yaml,
            "xlsx": self._process_xlsx,  # Add this line
            "xls": self._process_xlsx    # Add this line for older Excel files
        }
        
        processor = processors.get(file_type.lower())
        if not processor:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        return await processor(file_path, extraction_type)
    
    async def _process_docx(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process Word documents"""
        result = {
            "content": [],
            "metadata": {},
            "tables": [],
            "images": [],
            "structure": {}
        }

        try:
            doc = docx.Document(file_path)
            
            # Extract document properties
            result["metadata"] = {
                "core_properties": {
                    prop: getattr(doc.core_properties, prop)
                    for prop in dir(doc.core_properties)
                    if not prop.startswith('_')
                },
                "file_size": os.path.getsize(file_path),
                "filename": os.path.basename(file_path)
            }

            # Process paragraphs
            for para in doc.paragraphs:
                result["content"].append({
                    "text": para.text,
                    "style": para.style.name,
                    "runs": [{
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font": run.font.name if run.font.name else None
                    } for run in para.runs]
                })

            # Process tables
            if extraction_type in ["all", "tables"]:
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    result["tables"].append(table_data)

            # Add text analysis if content exists
            full_text = " ".join([p["text"] for p in result["content"]])
            if full_text and self.nlp:
                try:
                    doc = self.nlp(full_text)
                    result["analysis"] = {
                        "sentences": len(list(doc.sents)),
                        "words": len([token for token in doc if not token.is_punct]),
                        "entities": [
                            {
                                "text": ent.text,
                                "label": ent.label_,
                                "start": ent.start_char,
                                "end": ent.end_char
                            }
                            for ent in doc.ents
                        ]
                    }
                except Exception as e:
                    logger.warning(f"Text analysis failed: {e}")

            return result

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

    async def _process_csv(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process CSV files with enhanced analytics"""
        result = {
            "content": "",
            "metadata": {},
            "analysis": {},
            "preview": {},
            "statistics": {}
        }

        try:
            # Read first few rows to get column info and data types
            df_preview = pd.read_csv(file_path, nrows=5)
            total_rows = sum(1 for _ in open(file_path)) - 1  # Count total rows minus header

            # Get column info
            columns_info = {
                col: {
                    "dtype": str(df_preview[col].dtype),
                    "sample_values": df_preview[col].tolist()
                } for col in df_preview.columns
            }

            # Process in chunks for large files
            chunk_size = 10000
            chunks = pd.read_csv(file_path, chunksize=chunk_size)
            
            # Initialize statistics
            numeric_stats = {col: {
                "min": float('inf'),
                "max": float('-inf'),
                "sum": 0,
                "count": 0,
                "null_count": 0
            } for col in df_preview.columns if pd.api.types.is_numeric_dtype(df_preview[col])}

            categorical_stats = {col: {
                "unique_values": set(),
                "null_count": 0
            } for col in df_preview.columns if not pd.api.types.is_numeric_dtype(df_preview[col])}

            # Process each chunk
            preview_data = []
            processed_rows = 0
            for chunk in chunks:
                processed_rows += len(chunk)
                logger.info(f"Processing CSV: {processed_rows}/{total_rows} rows")

                # Store first 10 rows for preview
                if len(preview_data) < 10:
                    preview_data.extend(chunk.head(10 - len(preview_data)).to_dict('records'))

                # Update numeric statistics
                for col in numeric_stats.keys():
                    stats = numeric_stats[col]
                    col_data = chunk[col]
                    stats["min"] = min(stats["min"], col_data.min())
                    stats["max"] = max(stats["max"], col_data.max())
                    stats["sum"] += col_data.sum()
                    stats["count"] += col_data.count()
                    stats["null_count"] += col_data.isnull().sum()

                # Update categorical statistics
                for col in categorical_stats.keys():
                    stats = categorical_stats[col]
                    col_data = chunk[col]
                    stats["unique_values"].update(col_data.dropna().unique())
                    stats["null_count"] += col_data.isnull().sum()

            # Finalize numeric statistics
            for col, stats in numeric_stats.items():
                if stats["count"] > 0:
                    stats["mean"] = stats["sum"] / stats["count"]
                    stats["null_percentage"] = (stats["null_count"] / total_rows) * 100
                    del stats["sum"]  # Remove sum from final output
                    stats["unique_values"] = len(set(df_preview[col].dropna()))

            # Finalize categorical statistics
            for col, stats in categorical_stats.items():
                stats["unique_count"] = len(stats["unique_values"])
                stats["null_percentage"] = (stats["null_count"] / total_rows) * 100
                stats["unique_values"] = list(stats["unique_values"])[:10]  # Limit to 10 examples

            result.update({
                "content": str(preview_data),  # For vector store
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "total_rows": total_rows,
                    "total_columns": len(df_preview.columns),
                    "columns": list(df_preview.columns),
                    "columns_info": columns_info
                },
                "preview": {
                    "first_rows": preview_data,
                    "column_types": {col: str(dtype) for col, dtype in df_preview.dtypes.items()}
                },
                "statistics": {
                    "numeric_columns": numeric_stats,
                    "categorical_columns": categorical_stats
                },
                "analysis": {
                    "completeness": {
                        col: {
                            "filled": total_rows - (numeric_stats.get(col, {}).get("null_count", 0) or 
                                                  categorical_stats.get(col, {}).get("null_count", 0)),
                            "null_percentage": (numeric_stats.get(col, {}).get("null_percentage", 0) or 
                                              categorical_stats.get(col, {}).get("null_percentage", 0))
                        } for col in df_preview.columns
                    }
                }
            })

            return result

        except Exception as e:
            logger.error(f"Error processing CSV: {e}")
            raise
    
    async def _process_pdf(self, file_path: str, extraction_type: str) -> Dict[str, Any]:  # Make sure this is at same indentation level as __init__
        result = {
            "content": "",
            "metadata": {}
        }

        try:
            doc = fitz.open(file_path)
            result["metadata"] = doc.metadata

            text = ""
            for page in doc:
                text += page.get_text()
            
            result["content"] = text.strip()
            return result

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()



    async def _process_json(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
            """Process large JSON files"""
            result = {
                "content": "",
                "metadata": {},
                "analysis": {}
            }

            try:
                # Process JSON file in chunks using ijson for large files
                import ijson

                structure_sample = {}
                record_count = 0
                data_sample = []
                
                with open(file_path, 'rb') as file:
                    parser = ijson.parse(file)
                    
                    for prefix, event, value in parser:
                        if record_count < 1000:  # Analyze structure of first 1000 records
                            if prefix not in structure_sample:
                                structure_sample[prefix] = {
                                    "type": type(value).__name__,
                                    "sample": value if not isinstance(value, (dict, list)) else "complex"
                                }
                            
                            if prefix == "item" or prefix == "":  # Assuming array of objects or single object
                                if isinstance(value, dict) and record_count < 10:
                                    data_sample.append(value)
                                record_count += 1

                # Generate analysis
                result["content"] = str(data_sample)  # First 10 records as string for vector store
                result["metadata"] = {
                    "filename": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "record_count": record_count,
                    "structure": structure_sample
                }
                result["analysis"] = {
                    "preview_data": data_sample,
                    "schema": structure_sample
                }

                return result

            except Exception as e:
                logger.error(f"Error processing JSON: {e}")
                raise

    async def _process_xlsx(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process Excel files with enhanced analytics"""
        try:
            result = {
                "content": "",
                "metadata": {},
                "sheets": {},
                "analysis": {}
            }

            # Load workbook
            workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            
            all_sheets_data = {}
            total_cells = 0
            sheet_summaries = {}

            for sheet_name in workbook.sheetnames:
                try:
                    # Convert sheet to pandas DataFrame
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    if len(df) == 0:
                        continue  # Skip empty sheets

                    # Process sheet data
                    sheet_data = {
                        "headers": list(df.columns),
                        "preview_data": df.head(10).to_dict('records'),
                        "statistics": {
                            "numeric_columns": {},
                            "categorical_columns": {},
                            "row_count": len(df),
                            "column_count": len(df.columns)
                        }
                    }

                    # Calculate statistics for numeric columns
                    numeric_cols = df.select_dtypes(include=[np.number]).columns
                    for col in numeric_cols:
                        non_null_data = df[col].dropna()
                        stats = {
                            "min": float(non_null_data.min()) if not non_null_data.empty else None,
                            "max": float(non_null_data.max()) if not non_null_data.empty else None,
                            "mean": float(non_null_data.mean()) if not non_null_data.empty else None,
                            "null_count": int(df[col].isnull().sum()),
                            "null_percentage": float(df[col].isnull().sum() / len(df) * 100)
                        }
                        sheet_data["statistics"]["numeric_columns"][str(col)] = stats

                    # Calculate statistics for categorical columns
                    categorical_cols = df.select_dtypes(exclude=[np.number]).columns
                    for col in categorical_cols:
                        non_null_values = df[col].dropna()
                        value_counts = non_null_values.value_counts()
                        stats = {
                            "unique_count": len(value_counts),
                            "unique_values": value_counts.head(10).index.tolist(),  # First 10 unique values
                            "null_count": int(df[col].isnull().sum()),
                            "null_percentage": float(df[col].isnull().sum() / len(df) * 100)
                        }
                        sheet_data["statistics"]["categorical_columns"][str(col)] = stats

                    all_sheets_data[sheet_name] = sheet_data
                    total_cells += len(df) * len(df.columns)

                    # Create sheet summary
                    sheet_summaries[sheet_name] = {
                        "row_count": len(df),
                        "column_count": len(df.columns),
                        "numeric_columns": [str(col) for col in numeric_cols],
                        "categorical_columns": [str(col) for col in categorical_cols]
                    }

                except Exception as e:
                    logger.error(f"Error processing sheet {sheet_name}: {str(e)}")
                    continue

            # Create metadata
            result.update({
                "metadata": {
                    "filename": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "total_sheets": len(workbook.sheetnames),
                    "sheet_names": workbook.sheetnames,
                    "total_cells": total_cells,
                    "sheet_summaries": sheet_summaries
                },
                "sheets": all_sheets_data,
                "content": str({sheet: data["preview_data"] 
                              for sheet, data in all_sheets_data.items()}),
                "analysis": {
                    "sheets_overview": sheet_summaries,
                    "data_quality": {
                        sheet_name: {
                            "completeness": {
                                str(col): 100 - sheet_data["statistics"].get("numeric_columns", {}).get(str(col), {}).get("null_percentage", 0)
                                    if str(col) in sheet_data["statistics"].get("numeric_columns", {})
                                    else 100 - sheet_data["statistics"].get("categorical_columns", {}).get(str(col), {}).get("null_percentage", 0)
                                for col in sheet_data["headers"]
                            }
                        }
                        for sheet_name, sheet_data in all_sheets_data.items()
                    }
                }
            })

            return result

        except Exception as e:
            logger.error(f"Error processing Excel file: {e}")
            raise

        finally:
            if 'workbook' in locals():
                workbook.close()

    async def _process_xml(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def xml_to_dict(element):
                result = {}
                for child in element:
                    if len(child) == 0:
                        result[child.tag] = child.text
                    else:
                        result[child.tag] = xml_to_dict(child)
                return result

            data = xml_to_dict(root)
            
            return {
                "content": str(data),  # Convert to string for vector store
                "metadata": {
                    "root_tag": root.tag,
                    "size": os.path.getsize(file_path),
                    "filename": os.path.basename(file_path)
                },
                "structured_data": data
            }
        except Exception as e:
            logger.error(f"Error processing XML: {e}")
            raise

    async def _process_yaml(self, file_path: str, extraction_type: str) -> Dict[str, Any]:
        """Process YAML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = yaml.safe_load(file)
                return {
                    "content": str(data),  # Convert to string for vector store
                    "metadata": {
                        "type": "yaml",
                        "size": os.path.getsize(file_path),
                        "filename": os.path.basename(file_path)
                    },
                    "structured_data": data
                }
        except Exception as e:
            logger.error(f"Error processing YAML: {e}")
            raise
        